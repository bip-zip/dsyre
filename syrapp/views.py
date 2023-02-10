from django.shortcuts import render,redirect
from django.views import generic
from .models import Topic, Research
from django.utils.timezone import datetime
from django.http import HttpResponseRedirect
from .forms import EditorForm
from django.contrib import messages

class Home(generic.TemplateView):
    template_name = 'index.html'

    def get_context_data(self, **kwargs):
        context = super(Home, self).get_context_data(**kwargs) 
        today = datetime.today()
        td_topics=Topic.objects.filter(startdate__year=today.year, startdate__month=today.month, startdate__day=today.day)|Topic.objects.filter(enddate__year=today.year, enddate__month=today.month, enddate__day=today.day)
        history = Topic.objects.all().order_by('-startdate')
        unsettled = Topic.objects.filter(unsettled=True)
        context.update({'td_topics':td_topics,'history':history,'unsettled':unsettled,'home_active':'active'})
        return context
    
def update_topic(request):
    tid = request.POST['tid']
    data = request.POST['topic']
    topic =Topic.objects.get(id=tid)
    topic.topic = data
    topic.save()
    messages.success(request,'Topic title updated.')
    return redirect('/')

def add_topic(request):
    data = request.POST['topic']
    print(data)
    topic =Topic.objects.create(topic=data)
    topic.save()
    return redirect('/')

def add_research(request):
    tid = request.POST['topic']
    rid = request.POST['research']
    topicw = Topic.objects.get(id=tid)
    title = request.POST['title']
    link = request.POST['link']
    if Research.objects.filter(id=rid).first():
        research = Research.objects.get(id=rid)
        research.title=title
        research.link = link
        research.save()
    else:
        research =Research.objects.create(topic=topicw, title=title, link=link)
        research.save()
    return HttpResponseRedirect('/'+tid)



class TopicDetail(generic.DetailView):
    model = Topic
    template_name='topicdetail.html'


    def get_context_data(self, **kwargs):
        context = super(TopicDetail, self).get_context_data(**kwargs) 
        form = EditorForm()
        context.update({'detail_active':'active','form':form})
        return context



class UpdateResearch(generic.UpdateView):
    form_class=EditorForm
    model=Research
    template_name='updateresearch.html'

    def get_context_data(self, **kwargs):
        context = super(UpdateResearch, self).get_context_data(**kwargs) 
        rid  = self.kwargs['pk']
        research= Research.objects.get(id=rid)
        context.update({'detail_active':'active','research':research})
        return context
    
    def get_success_url(self):
        return '/update-research/{}'.format(self.kwargs['pk'])
        #    return redirect('/update-research/', self.kwargs['pk'])
        #    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

    
def delete_research(request,id):
    instance = Research.objects.get(id=id)
    instance.delete()
    messages.success(request,('"{}"'.format(instance.title)+" has removed."))
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

from docx import Document
import html2text
import os
def saveDocx(request,id):
    directory = os.path.join(locateFolder(),'Research Documents')
    topic = Topic.objects.get(id=id)
    print(topic.researches)
    document = Document()
    document.add_heading('Research Title: {}'.format(topic.topic), level=1)
    for i in topic.researches:
        document.add_heading(i.title, level=2)
        # document.add_paragraph(i.content)
        document.add_paragraph(html2text.html2text(i.content))

    document.save('{}/{}.docx'.format(directory,topic.topic))
    messages.success(request,('"{}.docx"'.format(topic.topic)+" has downloaded."))
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

def locateFolder():
    directory = "dSYRE"
    parent_dir = "C:/"
    path = os.path.join(parent_dir, directory)
    if not os.path.exists(path):
        os.mkdir(path)
        os.mkdir(os.path.join('C:/dSYRE', 'JupyterLab'))
        os.mkdir(os.path.join('C:/dSYRE', 'Research Documents'))
        

    print(path)
    return path



def openDF(request):
    # specify the directory path
    directory = locateFolder()
# open the directory
    os.startfile(directory)
    return redirect('/')

import subprocess
def openJL(request):
    directory =  os.path.join(locateFolder(),'JupyterLab')
    subprocess.run('start cmd /k "cd /d %cd% && deactivate myenv & C: & cd {} & conda activate & jupyter lab"'.format(directory), shell=True)
    return redirect('/')

def checkSettled(request,id,status):
    topic=Topic.objects.get(id=id)
    if status =='checked':
        topic.unsettled = False
        topic.save()
        messages.success(request,"👏 Seems like you have got what you need. Good going!")
    else:
        topic.unsettled = True
        topic.save()
        messages.success(request,"Seems like you want more. Keep going!")
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))



